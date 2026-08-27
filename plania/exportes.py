# © 2026 Martín Viera. Todos los derechos reservados.
"""
Plania · Exportes profesionales (PDF, Word, Excel)
==================================================
Convierte cualquier paquete de sugerencias o respuesta del copiloto en un
documento presentable para mandar al dueño, al comprador o al vendedor de la
ruta. Estructura común: título + secciones, donde cada sección tiene un
subtítulo, un texto explicativo ("el porqué") y una tabla con la evidencia.

    secciones = [(subtitulo, texto, DataFrame|None), ...]

Disponible en español, inglés y portugués vía el parámetro `idioma` de
`a_pdf`/`a_word`/`a_excel` — encabezado, pie de página y el separador de
miles de cada celda salen en el idioma pedido. Lo que NO cambia con
`idioma` es el contenido de `secciones`: quien arma la lista (`app/app.py`,
`plania/copiloto.py`) es quien decide en qué idioma va cada subtítulo y
texto — acá sólo se envuelve.
"""
from __future__ import annotations

import io
from datetime import date

import pandas as pd

from plania import i18n

AZUL = (31, 61, 122)      # marca Plania
GRIS = (245, 246, 248)

MAX_FILAS = 40
MAX_COLS = 8

# `plania/api.py` (la API que consume la ventana Electron/React, todavía
# 100% en español — es su propio cambio, aparte de este) lee `TITULOS` como
# un diccionario plano {clave: (titulo, texto)}. Se lo arma en español desde
# el catálogo para no tocarle el contrato a ese consumidor; `titulo_seccion`
# de acá abajo es la versión que sí recibe idioma, para app.py.
_CLAVES_SECCION = ("ofertas", "reposicion", "precios", "zonas", "recupero")


def titulo_seccion(clave: str, idioma: str = "es") -> tuple[str, str]:
    """(subtítulo, texto) de una sección estándar, en el idioma pedido."""
    return (i18n.t(f"exportes.titulo_{clave}", idioma),
            i18n.t(f"exportes.texto_{clave}", idioma))


TITULOS = {clave: titulo_seccion(clave, "es") for clave in _CLAVES_SECCION}


def _preparar(df: pd.DataFrame) -> pd.DataFrame:
    d = df.copy().head(MAX_FILAS)
    if len(d.columns) > MAX_COLS:
        d = d[[c for c in d.columns if c != "motivo"][:MAX_COLS]]
    for c in d.columns:
        if pd.api.types.is_float_dtype(d[c]):
            d[c] = d[c].round(2)
        if pd.api.types.is_datetime64_any_dtype(d[c]):
            d[c] = d[c].dt.date
    return d


def _celda(val, idioma: str) -> str:
    """Un valor de tabla como texto, con el separador de miles del idioma.

    Antes era `f"{val:,.2f}"` — el formato de Estados Unidos de Python,
    sin importar el idioma del documento. El resto del producto (tarjetas,
    tablas en pantalla, respuestas del copiloto) ya usa el separador de
    cada idioma; los exportes eran el único lugar que se había quedado
    mostrando `1,430,318.50` en un documento en español o portugués.
    """
    return i18n.miles(val, 2, idioma) if isinstance(val, float) else str(val)


# ---------------------------------------------------------------------------
# PDF (fpdf2)
# ---------------------------------------------------------------------------
def a_pdf(titulo: str, secciones: list, idioma: str = "es") -> bytes:
    from fpdf import FPDF

    pdf = FPDF(orientation="L", unit="mm", format="A4")
    pdf.set_auto_page_break(auto=True, margin=14)
    pdf.add_page()

    pdf.set_fill_color(*AZUL)
    pdf.rect(0, 0, 297, 26, style="F")
    pdf.set_text_color(255, 255, 255)
    pdf.set_font("Helvetica", "B", 17)
    pdf.set_xy(12, 7)
    pdf.cell(0, 8, _lat(f"Plania · {titulo}"))
    pdf.set_font("Helvetica", "", 10)
    pdf.set_xy(12, 16)
    pdf.cell(0, 6, _lat(i18n.t("exportes.encabezado_generado", idioma,
                               fecha=f"{date.today():%d/%m/%Y}")))
    pdf.set_y(32)

    for subtitulo, texto, df in secciones:
        pdf.set_text_color(*AZUL)
        pdf.set_font("Helvetica", "B", 13)
        pdf.cell(0, 9, _lat(subtitulo), new_x="LMARGIN", new_y="NEXT")
        if texto:
            pdf.set_text_color(40, 40, 40)
            pdf.set_font("Helvetica", "", 10)
            pdf.multi_cell(0, 5.5, _lat(texto))
            pdf.ln(1)
        if df is not None and len(df):
            _tabla_pdf(pdf, _preparar(df), idioma)
        pdf.ln(4)

    pdf.set_text_color(130, 130, 130)
    pdf.set_font("Helvetica", "I", 8)
    pdf.cell(0, 6, _lat(i18n.t("exportes.pie_marca", idioma)))
    return bytes(pdf.output())


def _tabla_pdf(pdf, d: pd.DataFrame, idioma: str = "es") -> None:
    ancho = 297 - 24
    w = ancho / len(d.columns)
    pdf.set_font("Helvetica", "B", 8)
    pdf.set_fill_color(*AZUL)
    pdf.set_text_color(255, 255, 255)
    for c in d.columns:
        pdf.cell(w, 6.5, _lat(str(c)[:24]), border=0, fill=True)
    pdf.ln()
    pdf.set_font("Helvetica", "", 8)
    pdf.set_text_color(30, 30, 30)
    for i, (_, fila) in enumerate(d.iterrows()):
        pdf.set_fill_color(*(GRIS if i % 2 == 0 else (255, 255, 255)))
        for val in fila:
            pdf.cell(w, 6, _lat(_celda(val, idioma)[:30]), border=0, fill=True)
        pdf.ln()


def _lat(s: str) -> str:
    return s.encode("latin-1", "replace").decode("latin-1")


# ---------------------------------------------------------------------------
# Word (python-docx)
# ---------------------------------------------------------------------------
def a_word(titulo: str, secciones: list, idioma: str = "es") -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor

    doc = Document()
    h = doc.add_heading(f"Plania · {titulo}", level=0)
    for run in h.runs:
        run.font.color.rgb = RGBColor(*AZUL)
    doc.add_paragraph(i18n.t("exportes.encabezado_generado", idioma,
                             fecha=f"{date.today():%d/%m/%Y}") + ".")

    for subtitulo, texto, df in secciones:
        doc.add_heading(subtitulo, level=1)
        if texto:
            doc.add_paragraph(texto)
        if df is not None and len(df):
            d = _preparar(df)
            tabla = doc.add_table(rows=1, cols=len(d.columns))
            tabla.style = "Light Grid Accent 1"
            for j, c in enumerate(d.columns):
                celda = tabla.rows[0].cells[j]
                celda.text = str(c)
                for p in celda.paragraphs:
                    for run in p.runs:
                        run.font.bold = True
                        run.font.size = Pt(9)
            for _, fila in d.iterrows():
                celdas = tabla.add_row().cells
                for j, val in enumerate(fila):
                    celdas[j].text = _celda(val, idioma)
                    for p in celdas[j].paragraphs:
                        for run in p.runs:
                            run.font.size = Pt(8)

    doc.add_paragraph().add_run(i18n.t("exportes.pie_marca_word", idioma)).italic = True
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Excel (xlsxwriter): una hoja por sección, con formato
# ---------------------------------------------------------------------------
def a_excel(secciones: list, idioma: str = "es") -> bytes:
    buf = io.BytesIO()
    with pd.ExcelWriter(buf, engine="xlsxwriter") as xw:
        wb = xw.book
        fmt_h = wb.add_format({"bold": True, "bg_color": "#1F3D7A",
                               "font_color": "white", "border": 1})
        for subtitulo, _texto, df in secciones:
            if df is None or not len(df):
                continue
            hoja = subtitulo[:31].replace("/", "-")
            d = df.copy().head(1000)
            for c in d.columns:
                if pd.api.types.is_datetime64_any_dtype(d[c]):
                    d[c] = d[c].dt.date
            d.to_excel(xw, sheet_name=hoja, index=False, startrow=1)
            ws = xw.sheets[hoja]
            for j, c in enumerate(d.columns):
                ws.write(1, j, str(c), fmt_h)
                ancho = max(10, min(38, int(d[c].astype(str).str.len().max() or 10) + 2))
                ws.set_column(j, j, ancho)
            ws.write(0, 0, i18n.t("exportes.hoja_marca", idioma, subtitulo=subtitulo,
                                  fecha=f"{date.today():%d/%m/%Y}"))
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Paquete de sugerencias -> secciones estándar
# ---------------------------------------------------------------------------
def secciones_desde_paquete(paquete: dict, incluir: list | None = None,
                            idioma: str = "es") -> list:
    res = paquete.get("resumen", {})
    secciones = [(i18n.t("exportes.resumen_titulo", idioma),
                  i18n.t("exportes.resumen_texto", idioma,
                        capital_liberable=i18n.miles(res.get("capital_liberable", 0), 0, idioma),
                        venta_en_riesgo=i18n.miles(res.get("venta_en_riesgo", 0), 0, idioma),
                        margen_extra_mensual=i18n.miles(res.get("margen_extra_mensual", 0), 0, idioma),
                        venta_potencial_zonas=i18n.miles(res.get("venta_potencial_zonas", 0), 0, idioma),
                        venta_recuperable=i18n.miles(res.get("venta_recuperable", 0), 0, idioma)),
                  None)]
    for clave in (incluir or list(_CLAVES_SECCION)):
        df = paquete.get(clave)
        if df is not None and len(df):
            sub, txt = titulo_seccion(clave, idioma)
            secciones.append((sub, txt, df))
    return secciones
